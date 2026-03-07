import sys
import csv
import os
from docx import Document

def get_csv_data(csv_path):
    data = {}
    try:
        with open(csv_path, "r", encoding ="utf-8-sig") as f:
            reader = csv.reader(f)

            row = list(reader)
            if len(row) >= 2:
                headers = row[0]
                values = row[1]
                data = dict(zip(headers, values))
            return data
    except Exception as e:
        print(f"Reading CSV Error : (e)")
        return{}

def update_csv_file(csv_path, data_dict):
    try:
        headers = list(data_dict.keys())
        values = list(data_dict.values())

        with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(headers) # Write Header
            writer.writerow(values)  # Write Data
        return True
    except Exception as e:
        print(f"Error updating CSV: {e}")
        return False

def replace_paragraphs(paragraph, replacement):
    for para in paragraph:
        for run in para.runs:
            for key, value in replacement.items():
                placeholder = f"${{{key}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

        full_text = "".join(run.text for run in para.runs)
        changed = False

        for key, value in replacement.items():
            placeholder = f"${{{key}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)
                changed = True

        if changed and para.runs:
            para.runs[0].text = full_text
            for run in para.runs[1:]:
                run.text =""



# a function for FastAPI implementation
# This is for the test,
# python replace.py template.docx data.csv output.docx
# This is like a CLI tool so added : def csv_to_docx(input_docx, input_csv, output_docx) for API
def csv_to_docx(input_docx, input_csv, output_docx):

    replacement = {}

    with open(input_csv, "r", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        rows = list(reader)

        if len(rows) >= 2:
            headers = rows[0]
            values = rows[1]
            replacement = dict(zip(headers, values))

    doc = Document(input_docx)

    replace_paragraphs(doc.paragraphs, replacement)

    doc.save(output_docx)

    return output_docx



def main(input_docx, input_csv, output_docx):

    replacement = {}
    try:
        with open(input_csv, "r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) >= 2 and row[0].strip():
                    replacement[row[0].strip()] = row[1].strip()

    except Exception as e:
        print(f"CSV Reading Error : {e}")
        sys.exit(1)

    try:
        doc = Document(input_docx)

    except Exception as e:
        print(f"Word File Oppening Error : {e}")
        sys.exit(1)

    replace_paragraphs(doc.paragraphs, replacement)

    doc.save(output_docx)
    print(f"Document Saved : {output_docx}")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("python replace.py input.docx input.csv output.docx")
        sys.exit(1)

    main(sys.argv[1], sys.argv[2], sys.argv[3])